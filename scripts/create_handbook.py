from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Netflix_Content_Analysis_Complete_Project_Handbook.docx"

RED = RGBColor(229, 9, 20)
INK = RGBColor(23, 23, 27)
GRAY = RGBColor(100, 100, 110)
LIGHT = "F4F5F7"
DARK = "17171B"
WHITE = RGBColor(255, 255, 255)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", size=10.5, color=INK, bold=False, italic=False,
             before=0, after=6, align=WD_ALIGN_PARAGRAPH.LEFT, style=None, font="Calibri"):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        set_run(r, size=size, color=color, bold=bold, italic=italic, font=font)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(text), size=10.2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text), size=10.2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run(r, size=18, color=INK, bold=True, font="Arial")
    elif level == 2:
        set_run(r, size=14, color=RED, bold=True, font="Arial")
    else:
        set_run(r, size=11.5, color=INK, bold=True, font="Arial")
    return p


def add_callout(doc, label, text, fill="F9E9EA"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(label.upper() + "  "), size=9, color=RED, bold=True, font="Arial")
    set_run(p.add_run(text), size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_fill(hdr[i], DARK)
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        set_run(p.add_run(header), size=9, color=WHITE, bold=True, font="Arial")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows) % 2 == 0:
                set_cell_fill(cells[i], LIGHT)
            p = cells[i].paragraphs[0]
            set_run(p.add_run(str(value)), size=8.7)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_code(doc, code, caption=None):
    if caption:
        add_para(doc, caption, size=9, color=GRAY, bold=True, after=3)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_fill(cell, "F1F2F4")
    set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line in code.splitlines():
        r = p.add_run(line + "\n")
        set_run(r, size=7.8, color=INK, font="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def page_break(doc):
    doc.add_page_break()


def section_intro(doc, title, purpose):
    page_break(doc)
    add_heading(doc, title, 1)
    add_para(doc, purpose, size=11, color=GRAY, italic=True, after=10)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.15
for name in ("Heading 1", "Heading 2", "Heading 3"):
    styles[name].font.name = "Arial"

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(hp.add_run("NETFLIX CONTENT ANALYSIS  |  COMPLETE PROJECT HANDBOOK"), size=8, color=GRAY, bold=True, font="Arial")

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(fp.add_run("Prepared for Kshitij  |  June 2026"), size=8, color=GRAY)

# Cover: editorial_cover pattern.
add_para(doc, "PROJECT HANDBOOK", size=10, color=RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=80, after=22)
add_para(doc, "Netflix Content Analysis", size=30, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, font="Arial")
add_para(doc, "Beginner-to-Advanced Technical, Analytical, and Interview Guide", size=15, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
add_para(doc, "A complete explanation of the dataset, preprocessing pipeline, dashboard architecture, visualizations, algorithms, code, deployment, limitations, learning path, and interview preparation.", size=11.5, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=65)
add_para(doc, "Prepared for", size=9, color=GRAY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
add_para(doc, "Kshitij", size=18, color=RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
add_para(doc, "Project state documented: June 18, 2026", size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
add_callout(doc, "Truthful scope", "This handbook clearly separates the working browser dashboard from future PRD features. The current system performs descriptive analytics and a simple illustrative extrapolation. It does not train Prophet, ARIMA, Isolation Forest, or any production machine-learning model.", "F4F5F7")

# Reading guide / TOC
page_break(doc)
add_heading(doc, "How to Use This Handbook", 1)
add_para(doc, "Start with Parts I-III if you are new. Use Parts IV-VI to understand and defend the implementation. Use Parts VII-IX for learning, interviews, viva, and future engineering.", after=10)
toc = [
    ("Part I", "Business and product foundations"),
    ("Part II", "Architecture, workflow, files, and technologies"),
    ("Part III", "Dataset, cleaning, preprocessing, engineering, and EDA"),
    ("Part IV", "Visualizations, algorithms, mathematics, statistics, and ML truth"),
    ("Part V", "Every file, every function, and important code walkthroughs"),
    ("Part VI", "Methodology, deployment, challenges, limitations, and roadmap"),
    ("Part VII", "What I must learn to fully understand this project"),
    ("Part VIII", "Interview preparation: 140 answered questions"),
    ("Part IX", "How to explain the project in 30 seconds to 5 minutes"),
]
add_table(doc, ["Part", "Purpose"], toc, [1.0, 5.5])

add_heading(doc, "Project Status at a Glance", 2)
status_rows = [
    ("Implemented", "Responsive dashboard, real titles/credits dataset, filters, search, five views, CSV export, CSV/JSON import, country map, contributor joins, descriptive metrics."),
    ("Illustrative only", "Five-year trend projection and narrative insight cards. These are not trained ML outputs."),
    ("PRD future scope", "Airflow/dbt medallion pipeline, warehouse, FastAPI, OAuth, Prophet/ARIMA, Isolation Forest, scheduled reports, monitoring, and cloud deployment."),
]
add_table(doc, ["Status", "Meaning"], status_rows, [1.35, 5.15])

# Part I
section_intro(doc, "Part I - Business and Product Foundations", "What the project is, why it exists, and how it creates value.")
add_heading(doc, "1. Executive Summary", 2)
add_para(doc, "Netflix Content Analysis is a browser-based analytical dashboard that turns a supplied catalogue dataset into understandable business views. It helps a user explore what content exists, when it was released, where it was produced, which genres and ratings dominate, and which actors and directors appear most often.")
add_para(doc, "The implemented application is intentionally lightweight: HTML provides structure, CSS provides the visual system, JavaScript performs in-browser data processing, Python prepares the source files, D3 renders geographic geometry, and TopoJSON supplies compressed country boundaries. No database or backend server is required for the current MVP.")
add_para(doc, "The bundled data contains 5,850 titles and 77,801 contributor credits. The preprocessing script joins titles to actors/directors and emits a browser-ready JavaScript catalogue. The dashboard can also import a different CSV or JSON file during the current browser session.")

add_heading(doc, "2. Project Objective", 2)
for text in [
    "Provide a single place to explore Netflix-like catalogue metadata.",
    "Reduce manual spreadsheet work by making filters and summaries interactive.",
    "Show temporal, genre, rating, geographic, runtime, cast, and director patterns.",
    "Allow a non-programmer to upload another compatible dataset and immediately reuse the dashboard.",
    "Create a foundation that could later connect to a real API, warehouse, and trained forecasting service.",
]:
    add_bullet(doc, text)

add_heading(doc, "3. Problem Statement", 2)
add_para(doc, "Raw CSV files are difficult for many stakeholders to use. A person must know formulas, pivot tables, joins, and data-cleaning rules before they can answer simple questions. Data is also easy to interpret inconsistently when each analyst creates a separate spreadsheet.")
add_callout(doc, "Core problem", "How can catalogue metadata and contributor credits be converted into a consistent, interactive, and understandable decision-support interface without requiring the user to write code?")

add_heading(doc, "4. Real-World Use Cases", 2)
use_cases = [
    ("Content strategist", "Compare genre concentration and release trends before discussing acquisition priorities."),
    ("Regional director", "Inspect production-country representation and compare regional catalogue mix."),
    ("Analyst", "Filter, search, and export a clean subset for a report or presentation."),
    ("Data scientist", "Understand available metadata and identify fields suitable for later modeling."),
    ("Student or interviewer", "Demonstrate ETL, frontend analytics, geospatial visualization, and honest model-scoping decisions."),
]
add_table(doc, ["User", "How the system helps"], use_cases, [1.45, 5.05])

# Part II
section_intro(doc, "Part II - Architecture and Technology", "How the system is organized from source data to browser output.")
add_heading(doc, "5. Project Architecture", 2)
add_code(doc, """Netflix dataset.zip
  |-- titles.csv -----------+
  |-- credits.csv ----------+--> Python preprocessing
                                  |-- clean fields
                                  |-- join contributors
                                  |-- derive region/genre
                                  v
                            data/catalogue.js
                                  |
index.html + styles.css + app.js + D3 + TopoJSON
                                  |
                                  v
                         Browser dashboard
                    filters | charts | map | export
                                  ^
                                  |
                       optional CSV/JSON import""", "Logical architecture")
add_para(doc, "This is a static client-side architecture. Static means the browser downloads files and runs the logic locally. Client-side means computation occurs on the user's machine rather than on a remote application server.")

add_heading(doc, "Architecture Layers", 3)
layers = [
    ("Source layer", "ZIP archive, titles.csv, credits.csv, Natural Earth world boundaries."),
    ("Preparation layer", "scripts/build_catalogue.py parses, cleans, joins, and serializes records."),
    ("Data layer", "catalogue.js stores normalized records; countries-110m.js stores map topology."),
    ("Presentation layer", "index.html defines controls and containers; styles.css controls appearance."),
    ("Analytics layer", "app.js filters, groups, ranks, calculates metrics, and creates SVG/HTML charts."),
    ("Runtime layer", "A simple Python HTTP server serves files to the browser."),
]
add_table(doc, ["Layer", "Responsibility"], layers, [1.35, 5.15])

add_heading(doc, "6. Folder Structure Explanation", 2)
folders = [
    ("index.html", "Application shell: sidebar, header, filters, modal, and script loading order."),
    ("styles.css", "All layout, colors, typography, responsive breakpoints, map, tables, and modal styling."),
    ("app.js", "Dashboard state, import logic, analytics, chart generation, map rendering, events, and export."),
    ("data/titles.csv", "5,850 title-level source rows."),
    ("data/credits.csv", "77,801 actor/director credit rows."),
    ("data/catalogue.js", "Generated normalized title records consumed directly by the browser."),
    ("data/countries-110m.*", "Natural Earth country boundaries in TopoJSON and JavaScript form."),
    ("scripts/build_catalogue.py", "Repeatable ETL/preprocessing script."),
    ("vendor/d3.min.js", "Local D3 library for projections, selections, CSV parsing, and SVG rendering."),
    ("vendor/topojson-client.min.js", "Converts compressed TopoJSON geometry to GeoJSON/SVG-ready shapes."),
    ("README.md", "Quick start, scope, file summary, and production boundaries."),
    ("Netflix_Content_Analysis_PRD.docx", "Original product requirements and long-term vision."),
]
add_table(doc, ["Path", "Why it exists"], folders, [2.05, 4.45])

add_heading(doc, "7. Complete Workflow", 2)
workflow = [
    "The source ZIP is copied into the workspace and extracted.",
    "Python reads titles.csv and credits.csv with UTF-8-aware CSV parsing.",
    "List-like text such as \"['drama', 'crime']\" is safely parsed into arrays.",
    "Credits are grouped by title ID; directors and a limited cast list are attached to each title.",
    "Country codes are converted to readable names and broad regions.",
    "The normalized records are written into data/catalogue.js.",
    "The browser loads libraries, world topology, catalogue data, and app.js in that order.",
    "app.js initializes state, populates filters, and renders the Overview view.",
    "Each interaction recomputes the active data subset and replaces the view HTML.",
    "CSV export serializes the active filtered subset.",
    "Optional import reads a user CSV/JSON file, recognizes common aliases, normalizes rows, and replaces the in-memory catalogue.",
]
for step in workflow:
    add_number(doc, step)

add_heading(doc, "24-25. Libraries and Technologies Used", 2)
tech = [
    ("HTML5", "Semantic page structure and native controls.", "Simple, accessible, universal.", "No analytics behavior by itself.", "React, Vue, Svelte."),
    ("CSS3", "Grid, flexbox, responsive layout, color system, modal styling.", "Fast and dependency-free.", "Large stylesheets require discipline.", "Tailwind, Bootstrap, Sass."),
    ("JavaScript ES2020+", "State, parsing, filtering, aggregation, rendering, events.", "Runs in every modern browser.", "Main-thread limits and weaker type safety.", "TypeScript, React, Web Workers."),
    ("Python 3", "Offline ETL and dataset generation.", "Excellent CSV/text ecosystem.", "Separate preprocessing step.", "Node.js, dbt, Spark."),
    ("python-docx", "Generates this handbook before PDF conversion.", "Programmatic Word layout.", "Complex layout requires OOXML care.", "ReportLab, LaTeX."),
    ("D3.js 7.9", "CSV parsing, geographic projection, SVG selection/rendering.", "Powerful and precise.", "Learning curve.", "Plotly, Vega-Lite, Leaflet."),
    ("TopoJSON Client 3.1", "Decompresses shared country boundaries.", "Smaller geographic files.", "Not beginner-friendly.", "GeoJSON, vector tiles."),
    ("Natural Earth / world-atlas", "Real-world country geometry at 110m resolution.", "Public, lightweight, appropriate for world maps.", "Not suitable for street detail.", "OpenStreetMap, Mapbox."),
    ("Python HTTP server", "Local static-file serving on port 8000.", "Zero configuration.", "Not production grade.", "Nginx, Vercel, Netlify."),
]
add_table(doc, ["Technology", "What/why", "Advantage", "Disadvantage", "Alternatives"], tech, [1.05, 1.75, 1.25, 1.25, 1.2])

# Part III
section_intro(doc, "Part III - Data and Analytics", "How the two source tables become analytical records.")
add_heading(doc, "8. Dataset Analysis", 2)
dataset_stats = [
    ("Titles", "5,850"),
    ("Credits", "77,801"),
    ("Movies", "3,744"),
    ("TV shows", "2,106"),
    ("Release-year range", "1945-2022"),
    ("Known directors", "4,041 titles"),
    ("Titles with cast", "5,340"),
    ("Production origins", "98 normalized values including Unknown"),
    ("Primary genres", "19"),
]
add_table(doc, ["Measure", "Observed value"], dataset_stats, [2.4, 4.1])
add_callout(doc, "Dataset caveat", "The dataset does not include Netflix date-added history, true regional availability, viewing hours, subscribers, budgets, or acquisition costs. Therefore the dashboard analyzes release year and production origin, not viewership performance.")

add_heading(doc, "9. Explanation of Every Source Column", 2)
title_columns = [
    ("id", "Unique title identifier; join key to credits.csv."),
    ("title", "Human-readable movie or show name."),
    ("type", "MOVIE or SHOW; normalized to Movie or TV Show."),
    ("description", "Synopsis text; preserved for future search or NLP."),
    ("release_year", "Year of original release; primary temporal field."),
    ("age_certification", "Audience maturity label such as TV-MA or PG-13."),
    ("runtime", "Minutes per movie or typical episode, depending on source semantics."),
    ("genres", "String representation of a list of genre labels."),
    ("production_countries", "String representation of ISO-like country-code list."),
    ("seasons", "Number of seasons for shows; blank for movies."),
    ("imdb_id", "External IMDb identifier."),
    ("imdb_score", "Average IMDb user score, normally 0-10."),
    ("imdb_votes", "Count of IMDb votes; useful for confidence/popularity context."),
    ("tmdb_popularity", "TMDb popularity metric."),
    ("tmdb_score", "Average TMDb score."),
]
add_table(doc, ["titles.csv column", "Meaning and use"], title_columns, [1.65, 4.85])
credit_columns = [
    ("person_id", "Unique contributor identifier."),
    ("id", "Title ID used to join to titles.csv."),
    ("name", "Actor or director name."),
    ("character", "Character played; mainly relevant for actors."),
    ("role", "ACTOR or DIRECTOR; controls which normalized array receives the person."),
]
add_table(doc, ["credits.csv column", "Meaning and use"], credit_columns, [1.65, 4.85])

add_heading(doc, "Normalized Browser Record", 3)
normalized = [
    ("genre", "First genre, used as the primary grouping dimension."),
    ("genres", "Full genre array retained for future multi-label analysis."),
    ("country", "Readable primary production country."),
    ("countryCode", "Primary source country code."),
    ("region", "Derived broad business region."),
    ("duration", "Numeric runtime."),
    ("seasons", "Numeric season count."),
    ("director", "Up to three directors joined as a string."),
    ("cast", "Up to eight actors joined as an array."),
    ("imdbScore / imdbVotes", "Numeric score and vote count."),
]
add_table(doc, ["Derived field", "Purpose"], normalized, [1.7, 4.8])

add_heading(doc, "10. Data Cleaning Process", 2)
cleaning = [
    ("Encoding", "Files are opened as UTF-8 with BOM handling to avoid broken punctuation and names."),
    ("Missing values", "Missing genre becomes Unclassified; country/director become Unknown; ratings become Not rated; numeric fields default safely."),
    ("List parsing", "Python ast.literal_eval converts list-like strings without executing arbitrary code."),
    ("Type normalization", "MOVIE/SHOW become consistent display labels."),
    ("Country normalization", "Known codes are mapped to names; unknown codes are preserved."),
    ("Numeric conversion", "Runtime, seasons, scores, and votes become int/float or null/default values."),
    ("Join control", "Credits are grouped by title ID; only valid role values are assigned."),
    ("Browser import validation", "Imported rows require a title and plausible year between 1800 and 2100."),
]
add_table(doc, ["Cleaning task", "Implementation"], cleaning, [1.65, 4.85])

add_heading(doc, "11. Data Preprocessing", 2)
add_para(doc, "Preprocessing means converting raw data into a consistent representation before analysis. The Python script builds a dictionary keyed by title ID, scans credits once, then scans titles once. This avoids repeatedly searching all 77,801 credits for every title.")
add_code(doc, """credits = defaultdict(lambda: {"directors": [], "cast": []})
for row in credits_csv:
    target = credits[row["id"]]
    if row["role"] == "DIRECTOR":
        target["directors"].append(row["name"])
    elif row["role"] == "ACTOR":
        target["cast"].append(row["name"])""", "Core join preparation")
add_para(doc, "Time complexity is approximately O(C + T), where C is the number of credit rows and T is the number of title rows. A naive nested loop would be O(C × T), which would be much slower.")

add_heading(doc, "12. Feature Engineering", 2)
features = [
    ("Primary genre", "Uses the first genre for simple mutually exclusive grouping. Advantage: easy charts. Disadvantage: loses multi-genre nuance."),
    ("Primary country", "Uses the first production country. Advantage: simple map/ranking. Disadvantage: co-productions are underrepresented."),
    ("Region", "Maps a country to North America, LATAM, APAC, Africa, or Europe/Other."),
    ("Contributor counts", "Counts appearances per director and actor."),
    ("Collaboration pairs", "Creates actor × director keys and counts repeat combinations."),
    ("Runtime buckets", "Groups movie runtimes into interpretable ranges."),
    ("Recent release flag", "Counts titles released from 2020 onward."),
    ("Illustrative projection", "Extrapolates recent annual volume using a deterministic trend formula; not a trained feature/model."),
]
add_table(doc, ["Engineered feature", "Explanation and trade-off"], features, [1.7, 4.8])

add_heading(doc, "13. Exploratory Data Analysis", 2)
for text in [
    "Univariate analysis: counts of type, genre, rating, country, year, and runtime bucket.",
    "Bivariate analysis: genre by region, actor-director pair frequency, type proportions, and country rankings.",
    "Temporal analysis: annual title-volume trend over the most recent years.",
    "Missingness awareness: Unknown and Not rated remain visible rather than silently disappearing.",
    "Interactive slicing: the same analysis is recalculated after filters and search.",
]:
    add_bullet(doc, text)
add_para(doc, "EDA is descriptive. It reveals patterns and data-quality issues but does not prove causality. For example, many US titles do not prove that US content causes subscriber growth.")

# Part IV
section_intro(doc, "Part IV - Visualizations, Algorithms, Mathematics, and ML", "What calculations are used, what is not used, and why.")
add_heading(doc, "14. Visualizations Used", 2)
visuals = [
    ("KPI cards", "Immediate totals and proportions.", "Fast scanning.", "Can oversimplify and currently contains some illustrative comparison text."),
    ("Line/area chart", "Annual release volume and illustrative outlook.", "Shows direction over time.", "Area can visually exaggerate magnitude."),
    ("Horizontal bars", "Genre and market rankings.", "Excellent category comparison.", "Long labels need space."),
    ("Donut chart", "Rating/region composition.", "Compact share overview.", "Harder than bars for precise comparison."),
    ("Vertical bars", "Runtime bucket distribution.", "Good discrete comparison.", "Bucket boundaries influence interpretation."),
    ("Heatmap", "Genre concentration by region.", "Finds clusters quickly.", "Color perception and scale need care."),
    ("World map", "Production-origin footprint.", "Geographic context.", "Marker placement covers selected major markets, not all 98 origins."),
    ("Data table", "Exact title-level records.", "Precise lookup/export context.", "Wide tables can require horizontal scrolling."),
]
add_table(doc, ["Visualization", "Purpose", "Advantage", "Disadvantage"], visuals, [1.25, 2.0, 1.55, 1.7])

add_heading(doc, "15-17. Algorithms Used, Reasons, and Alternatives", 2)
algorithms = [
    ("Filtering", "Array.filter tests type, region, genre, year, and search predicates.", "Simple and correct for 5,850 rows.", "Indexed database queries for millions of rows."),
    ("Grouping/counting", "Array.reduce builds frequency dictionaries.", "One-pass O(n) aggregation.", "SQL GROUP BY, d3.rollup, Map."),
    ("Ranking", "Object.entries + sort descending + slice.", "Readable top-N logic.", "Heap/partial selection for huge data."),
    ("Join", "Python dictionary keyed by title ID.", "Near-linear join behavior.", "pandas merge, SQL JOIN, Spark join."),
    ("CSV parsing", "Python csv module offline; D3 csvParse for browser imports.", "Handles quoted commas and headers.", "Papa Parse for streaming/large files."),
    ("List parsing", "ast.literal_eval offline; conservative JSON/comma parser in browser.", "Safer than eval.", "Strict JSON source format."),
    ("Projection", "D3 Natural Earth projection converts longitude/latitude to SVG x/y.", "Accurate global display.", "Mercator, Robinson, Equal Earth."),
    ("TopoJSON conversion", "Shared arcs are expanded to GeoJSON features/mesh.", "Small file and clean borders.", "Raw GeoJSON or vector tiles."),
    ("Illustrative forecast", "Recent value × deterministic growth plus sinusoidal adjustment.", "Demonstrates UI shape without backend.", "Prophet, ARIMA, ETS, regression; required for real forecasting."),
]
add_table(doc, ["Algorithm", "How it works", "Why chosen", "Alternative"], algorithms, [1.25, 2.15, 1.55, 1.55])

add_heading(doc, "18-22. Model Training, Tuning, Evaluation, Interpretation, Importance", 2)
add_callout(doc, "Critical interview answer", "No machine-learning model is trained in the current implementation. Therefore there is no train/test split, fitted parameter set, hyperparameter search, MAPE result, feature importance output, or production model artifact.")
add_heading(doc, "What a Real Forecasting Pipeline Would Do", 3)
ml_steps = [
    "Aggregate genuine monthly Netflix addition dates by genre.",
    "Order observations by time and reserve the last 12 months as a holdout set.",
    "Fit a baseline seasonal-naive model first.",
    "Fit Prophet, ARIMA/SARIMA, exponential smoothing, or gradient-boosted lag models.",
    "Tune changepoint prior, seasonality modes, ARIMA p/d/q and seasonal P/D/Q using time-series cross-validation.",
    "Evaluate MAE, RMSE, MAPE/SMAPE, interval coverage, and performance against the baseline.",
    "Inspect residual autocorrelation and failure periods.",
    "Register the selected model, version its data, and monitor drift after deployment.",
]
for step in ml_steps:
    add_number(doc, step)
add_heading(doc, "Hyperparameter Examples", 3)
add_table(doc, ["Model", "Important hyperparameters", "Meaning"], [
    ("Prophet", "changepoint_prior_scale, seasonality_prior_scale, seasonality_mode", "Trend flexibility and seasonal strength/form."),
    ("SARIMA", "p,d,q,P,D,Q,s", "Autoregression, differencing, moving average, and seasonal structure."),
    ("Isolation Forest", "n_estimators, contamination, max_samples", "Tree count, expected anomaly rate, sample size."),
    ("Gradient boosting", "learning_rate, max_depth, n_estimators", "Step size, complexity, and number of trees."),
], [1.1, 2.8, 2.6])
add_heading(doc, "Feature Importance", 3)
add_para(doc, "The current descriptive dashboard has no model feature importance. In a future tree model, permutation importance or SHAP could estimate how lagged volume, genre, month, country, rating mix, and seasonality influence predictions. Importance is association with model output, not proof of causation.")

add_heading(doc, "29-31. Mathematics, Statistics, and Machine-Learning Concepts", 2)
concept_rows = [
    ("Count", "Number of records satisfying a condition: n = Σ I(condition)."),
    ("Percentage", "part / whole × 100; used for movie/show and rating shares."),
    ("Mean", "Σx / n; suitable for average score when missing values are excluded."),
    ("Frequency distribution", "Count for each category; foundation of bar and donut charts."),
    ("Ranking", "Sort categories by count to identify top-N values."),
    ("Binning", "Convert continuous runtime into intervals; improves readability but loses detail."),
    ("Time series", "Observations indexed by time; order matters and random train/test split is usually wrong."),
    ("Growth rate", "(new - old) / old × 100; must not be claimed when comparison data is absent."),
    ("MAE", "Mean absolute error: average |actual - forecast|."),
    ("MAPE", "Mean absolute percentage error; unstable when actual values are zero."),
    ("Confidence/prediction interval", "Range intended to contain an unknown estimate/future observation at a stated rate."),
    ("Cosine similarity", "Angle-based vector similarity proposed in the PRD for genre mix comparison."),
    ("Anomaly detection", "Finding observations that differ strongly from normal patterns."),
    ("Bias/variance", "Underfitting versus sensitivity to training data."),
    ("Data leakage", "Using future or target information during training, producing unrealistically good results."),
]
add_table(doc, ["Concept", "Simple explanation"], concept_rows, [1.55, 4.95])

# Part V
section_intro(doc, "Part V - Files, Functions, and Code Walkthrough", "A practical guide to what the implementation does.")
add_heading(doc, "26. Explanation of Every File", 2)
for path, meaning in folders:
    add_heading(doc, path, 3)
    add_para(doc, meaning)

add_heading(doc, "27. Explanation of Every JavaScript Function", 2)
function_details = {
    "firstValue": "Finds the first non-empty value matching any accepted column alias. This makes imports tolerant of different header names.",
    "parseListValue": "Converts arrays, JSON-like strings, or comma/pipe/semicolon text into a clean JavaScript array.",
    "normalizeType": "Maps show/series variants to TV Show and all other values to Movie.",
    "normalizeCountry": "Reads the first country and converts common country codes into readable names.",
    "inferRegion": "Returns an explicit region when supplied; otherwise maps known countries to broad regions.",
    "numberValue": "Removes non-numeric characters and safely parses a number with a fallback.",
    "normalizeRecord": "Creates the canonical title object used by every dashboard view and rejects rows without a valid title/year.",
    "resetDashboardState": "Clears filters/search and rebuilds filter options after a dataset change.",
    "updateDatasetStatus": "Updates the header with the active filename and latest release year.",
    "importDataset": "Reads CSV/JSON, parses rows, normalizes records, reports skipped rows, swaps the active catalogue, and rerenders.",
    "openImportModal": "Displays the import dialog.",
    "closeImportModal": "Hides the dialog and clears the native file input.",
    "filteredData": "Applies every active filter and the global text search to the catalogue.",
    "populateFilters": "Builds unique region, genre, and year options from the active dataset.",
    "metricCard": "Returns reusable KPI-card HTML.",
    "panel": "Returns reusable dashboard-panel HTML with title, subtitle, content, width, and optional action.",
    "lineChart": "Calculates SVG coordinates, grid lines, labels, area fill, line path, and tooltip hit areas.",
    "barChart": "Creates horizontal HTML bars or vertical SVG bars depending on the requested orientation.",
    "donutChart": "Calculates cumulative percentages and builds a CSS conic-gradient plus legend.",
    "overviewView": "Calculates type, recent-release, genre, rating, and annual metrics for the overview.",
    "trendView": "Builds release-year history, runtime buckets, recent genre ranking, and illustrative projection.",
    "geographyView": "Calculates country/region counts, market ranking, map panel, donut, and regional heatmap.",
    "worldMap": "Returns the empty accessible container that D3 later fills.",
    "renderWorldMap": "Converts TopoJSON to country paths, fits a Natural Earth projection, and draws borders and markers.",
    "regionHeatmap": "Counts top genres inside each configured region and converts relative intensity to red opacity.",
    "peopleView": "Counts directors, actors, and actor-director pairings for contributor intelligence.",
    "peopleCards": "Formats contributor leaderboard entries.",
    "libraryView": "Sorts and displays up to 100 matching titles and catalogue summary cards.",
    "libraryTable": "Creates the title/type/genre/origin/year/rating table.",
    "emptyState": "Shows a helpful message when filters return zero records.",
    "render": "Central controller: reads state, calls the active view, inserts HTML, binds tooltips, and triggers map rendering.",
    "bindTooltips": "Adds mouse listeners to elements with data-tip attributes.",
    "exportCsv": "Escapes filtered values, builds CSV text, creates a Blob URL, and triggers a download.",
    "showToast": "Displays a temporary status notification.",
}
for name, desc in function_details.items():
    add_heading(doc, name + "()", 3)
    add_para(doc, desc)
    if name in {"filteredData", "render", "normalizeRecord", "importDataset", "renderWorldMap"}:
        add_para(doc, "Why this matters: it is a central integration point. A defect here can affect multiple views, so it deserves focused tests.", size=9.5, color=GRAY, italic=True)

add_heading(doc, "Python Preprocessing Functions", 2)
py_funcs = [
    ("parse_list(value)", "Safely parses Python-list-looking text into a list; returns an empty list on malformed input."),
    ("region_for(code)", "Maps a country code into a broad region."),
    ("title_case_genre(value)", "Replaces underscores and title-cases a genre label."),
    ("Main credits loop", "Builds an ID-indexed contributor lookup."),
    ("Main titles loop", "Normalizes each title and attaches contributor information."),
    ("JSON serialization", "Writes compact UTF-8 JavaScript data into catalogue.js."),
]
add_table(doc, ["Python unit", "Purpose"], py_funcs, [1.8, 4.7])

add_heading(doc, "28 & WHAT EACH LINE OF CODE IS DOING", 2)
walkthroughs = [
    ("State and catalogue",
"""const bundledCatalogue = window.NETFLIX_CATALOGUE || [];
let catalogue = [...bundledCatalogue];
let datasetName = "Bundled dataset";""",
["Read the generated catalogue from the global browser object.", "Create a mutable copy so imported data can replace it.", "Store a human-readable name for the active data source."]),
    ("Filtering",
"""return catalogue.filter(item =>
  (state.type === "All" || item.type === state.type) &&
  (state.region === "All" || item.region === state.region) &&
  (state.genre === "All" || item.genre === state.genre) &&
  (state.year === "All" || item.year === Number(state.year))
);""",
["filter visits every title.", "Each condition accepts all values when its filter is All.", "Otherwise the record must equal the selected value.", "&& means every active condition must be true."]),
    ("Grouping",
"""const group = (items, key) => items.reduce((acc, item) => {
  const value = typeof key === "function" ? key(item) : item[key];
  acc[value] = (acc[value] || 0) + 1;
  return acc;
}, {});""",
["reduce carries one accumulator object through the array.", "The grouping key can be a property name or a function.", "Missing counts start at zero, then increment.", "The final object maps category to frequency."]),
    ("Import normalization",
"""const normalized = rows.map(normalizeRecord).filter(Boolean);
if (!normalized.length) throw new Error("No valid rows found.");
catalogue = normalized;
resetDashboardState();
render();""",
["map converts each raw row into the canonical schema.", "filter(Boolean) removes rejected null rows.", "The guard prevents replacing the dashboard with empty invalid data.", "The active catalogue is replaced, controls reset, and every view rerenders."]),
    ("CSV export",
"""const blob = new Blob([csv], { type: "text/csv;charset=utf-8" });
const link = document.createElement("a");
link.href = URL.createObjectURL(blob);
link.download = filename;
link.click();""",
["Blob creates an in-memory file.", "A temporary anchor element is created.", "A temporary URL points to the in-memory bytes.", "download supplies the filename.", "click starts the browser download."]),
    ("Map projection",
"""const projection = d3.geoNaturalEarth1()
  .fitExtent([[18, 18], [width - 18, height - 18]], { type: "Sphere" });
const path = d3.geoPath(projection);""",
["Create a Natural Earth projection.", "fitExtent scales and centers the globe inside the SVG margins.", "geoPath converts geographic geometry into SVG path commands."]),
    ("Python join",
"""target = credits[row["id"]]
if row["role"] == "DIRECTOR":
    target["directors"].append(row["name"])
elif row["role"] == "ACTOR":
    target["cast"].append(row["name"])""",
["Use title ID as the dictionary key.", "Route directors into the director list.", "Route actors into the cast list.", "The prebuilt lookup is later attached to title records."]),
]
for title, code, explanations in walkthroughs:
    add_heading(doc, title, 3)
    add_code(doc, code)
    for i, explanation in enumerate(explanations, 1):
        add_para(doc, f"Line/step {i}: {explanation}", size=9.6, after=3)

# Part VI
section_intro(doc, "Part VI - Methodology, Deployment, Challenges, and Roadmap", "How the work was conducted and what should happen next.")
add_heading(doc, "32. Methodology (Research-Paper Style)", 2)
method_sections = [
    ("Research design", "Applied descriptive analytics and software prototyping. The goal was to construct an interactive artefact and evaluate whether supplied metadata could support useful catalogue exploration."),
    ("Data source", "A supplied archive containing titles.csv and credits.csv. No subscriber, viewership, or confidential Netflix operational data was used."),
    ("Population and sample", "The working population is the full supplied snapshot: 5,850 titles and 77,801 credits. It is not necessarily the complete historical Netflix catalogue."),
    ("Data preparation", "Schema inspection, encoding-safe parsing, missing-value handling, type conversion, country/region mapping, contributor grouping, and title-credit join."),
    ("Analysis", "Frequency analysis, top-N ranking, percentages, runtime binning, release-year trend inspection, regional cross-tabulation, and collaboration counting."),
    ("System development", "Iterative prototyping from PRD to dependency-free dashboard, followed by real-data integration, layout correction, real-world map replacement, and dataset importer."),
    ("Validation", "JavaScript syntax checks, record-count checks, HTTP 200 checks for application/data assets, and manual screenshot-driven UI correction."),
    ("Ethics and validity", "Public metadata only; unknown values are preserved; descriptive associations are not represented as causal effects."),
]
for title, text in method_sections:
    add_heading(doc, title, 3)
    add_para(doc, text)

add_heading(doc, "23. Deployment Workflow", 2)
deployment = [
    ("Current local", "Run python -m http.server 8000 in the project folder, then open http://localhost:8000."),
    ("Static hosting", "Upload HTML/CSS/JS/data/vendor files to GitHub Pages, Netlify, Vercel static hosting, S3+CloudFront, or Nginx."),
    ("Production API", "Move large data and calculations behind FastAPI; add pagination, caching, authentication, logging, and rate limiting."),
    ("Data pipeline", "Schedule ingestion/quality/transformation with Airflow and dbt; store curated data in a warehouse."),
    ("ML service", "Train/version models separately, expose forecasts via API, and monitor accuracy/drift."),
    ("CI/CD", "Lint/test/build on every commit, deploy preview environments, then promote approved releases."),
]
add_table(doc, ["Stage", "Workflow"], deployment, [1.25, 5.25])

add_heading(doc, "33. Challenges Faced During Development", 2)
challenges = [
    ("Unicode extraction", "PowerShell's legacy encoding could not print a PRD symbol. UTF-8 output was explicitly enabled."),
    ("Workspace permissions", "Files in Downloads required approved copying into the writable workspace."),
    ("Dataset mismatch", "The supplied dataset lacked date_added and availability; labels and analysis were changed to release year and production origin."),
    ("Wide-card layout", "A three-column strategy card became unreadable. It was removed and the table width expanded."),
    ("Map quality", "A rough placeholder map was replaced with Natural Earth/TopoJSON geometry."),
    ("Map fetch failure", "Browser fetch restrictions caused map loading to fail. Topology was packaged as a local JavaScript global."),
    ("Server lifetime", "A temporary smoke-test server ended, causing connection refused. A persistent background server was started."),
    ("Browser automation", "The in-app browser control process was blocked by Windows sandbox permissions, so structural/server tests and user screenshots guided corrections."),
]
add_table(doc, ["Challenge", "Resolution"], challenges, [1.45, 5.05])

add_heading(doc, "34. Limitations of the Current System", 2)
limitations = [
    "No real Netflix viewership, engagement, subscriber, licensing, or regional-availability data.",
    "Dataset ends in 2022 and may be incomplete or biased.",
    "Only the first genre and first country are used for several charts.",
    "Map markers cover selected major countries rather than every origin.",
    "The current forecast is illustrative and should not be used for decisions.",
    "Some KPI comparison phrases are design placeholders rather than calculated historical comparisons.",
    "All analytics run on the browser main thread; very large imports may become slow.",
    "Imported data is session-only and is not saved to a database.",
    "No authentication, authorization, audit log, automated tests, or accessibility certification.",
    "HTML strings include dataset values; production code should escape untrusted text to prevent injection.",
]
for item in limitations:
    add_bullet(doc, item)

add_heading(doc, "35. Future Improvements", 2)
future = [
    ("Immediate", "Calculate every KPI comparison from data; escape imported text; add importer preview/mapping; add unit tests; support all countries on the map."),
    ("Near term", "Move to TypeScript/React, add state management, virtualized tables, web workers, and saved filter presets."),
    ("Data", "Add genuine date-added, availability, viewership, rights, cost, and competitor data with governance."),
    ("Backend", "FastAPI + warehouse + Redis; pagination; OAuth/SSO; scheduled exports."),
    ("ML", "Real time-series backtesting, seasonal baselines, Prophet/SARIMA/ETS comparison, anomaly detection, model cards."),
    ("Operations", "Airflow/dbt, data-quality tests, observability, CI/CD, load testing, WCAG 2.1 AA."),
]
add_table(doc, ["Horizon", "Improvement"], future, [1.2, 5.3])

# Part VII
section_intro(doc, "WHAT I MUST LEARN TO FULLY UNDERSTAND THIS PROJECT", "A recommended beginner-to-advanced learning path.")
learning = [
    ("Beginner 1", "Computer/file basics", "Folders, extensions, ZIP/CSV/JSON, terminal, localhost, ports."),
    ("Beginner 2", "HTML", "Elements, attributes, forms, semantic structure, accessibility."),
    ("Beginner 3", "CSS", "Box model, selectors, flexbox, grid, responsive media queries."),
    ("Beginner 4", "JavaScript fundamentals", "Variables, arrays, objects, functions, conditions, loops, DOM events."),
    ("Beginner 5", "Data basics", "Rows, columns, data types, missing values, primary/foreign keys."),
    ("Intermediate 1", "Modern JavaScript", "map/filter/reduce, template strings, async/await, modules, Blob/File APIs."),
    ("Intermediate 2", "Data analysis", "Grouping, percentages, distributions, outliers, binning, EDA."),
    ("Intermediate 3", "Python", "CSV reading, dictionaries, defaultdict, functions, exceptions, JSON."),
    ("Intermediate 4", "Visualization", "Chart choice, scales, SVG, D3 selections, color and accessibility."),
    ("Intermediate 5", "Geospatial basics", "Longitude/latitude, projections, GeoJSON, TopoJSON, choropleths."),
    ("Intermediate 6", "Software engineering", "Git, testing, linting, separation of concerns, security."),
    ("Advanced 1", "Databases and SQL", "JOIN, GROUP BY, indexes, query plans, warehouses."),
    ("Advanced 2", "Data engineering", "ETL/ELT, Airflow, dbt, medallion layers, data quality."),
    ("Advanced 3", "Statistics", "Sampling, uncertainty, confidence intervals, hypothesis tests, causality."),
    ("Advanced 4", "Time-series ML", "Stationarity, lags, seasonality, ARIMA, Prophet, backtesting."),
    ("Advanced 5", "MLOps", "Experiment tracking, model registry, drift, monitoring, reproducibility."),
    ("Advanced 6", "Production architecture", "APIs, OAuth, caching, queues, cloud, CI/CD, observability."),
]
add_table(doc, ["Order", "Topic", "What to learn"], learning, [1.0, 1.7, 3.8])
add_callout(doc, "Recommended order", "Do not start with machine learning. First become comfortable reading the HTML/CSS/JavaScript and explaining the dataset. Then learn analytics and Python. Add databases and real forecasting only after those foundations.")

# Part VIII: generated Q&A
section_intro(doc, "INTERVIEW PREPARATION", "140 questions with answers, organized by difficulty and purpose.")

concepts = [
    ("HTML", "the language that gives a web page structure", "It defines controls, headings, navigation, and containers in index.html."),
    ("CSS", "the language that controls visual presentation", "It creates the grid, responsive design, colors, spacing, modal, and map styling."),
    ("JavaScript", "the programming language executed by the browser", "It performs filtering, calculations, rendering, import, export, and interaction."),
    ("CSV", "a text format where rows are records and commas separate fields", "It is used for titles.csv, credits.csv, imports, and exports."),
    ("JSON", "a structured text format based on objects and arrays", "It represents normalized records and imported data."),
    ("DOM", "the browser's object representation of HTML", "The app queries elements, changes HTML, and attaches event listeners."),
    ("Array", "an ordered collection of values", "The catalogue, cast lists, countries, and chart entries are arrays."),
    ("Object", "a collection of named key/value properties", "Each title and the global state are objects."),
    ("Function", "a reusable block of logic", "Functions isolate parsing, analytics, charts, views, and events."),
    ("State", "the current values that control the interface", "state stores the active view, filters, and search query."),
    ("Filter", "an operation that keeps items satisfying a condition", "filteredData creates the active analytical subset."),
    ("Reduce", "an operation that combines many array values into one result", "group uses reduce to build frequency counts."),
    ("ETL", "extract, transform, load", "Python extracts CSVs, transforms rows, and loads catalogue.js."),
    ("Primary key", "a unique identifier for a record", "Title id identifies a title."),
    ("Foreign key", "a value referencing another table's primary key", "credits.csv id points to titles.csv id."),
    ("Join", "combining related rows across tables", "Credits are attached to titles using id."),
    ("Missing value", "a field with no known value", "The pipeline uses Unknown, Not rated, zero, or null depending on meaning."),
    ("Normalization", "making inconsistent values follow one schema", "MOVIE/SHOW and country/genre formats are standardized."),
    ("Aggregation", "summarizing many rows", "Counts by genre, year, rating, and country are aggregations."),
    ("Visualization", "a graphical encoding of data", "Charts and maps make patterns easier to perceive."),
    ("SVG", "a vector graphics format represented by DOM elements", "Line charts and geographic paths are drawn with SVG."),
    ("D3", "a JavaScript library for data-driven documents", "It parses CSV and renders the Natural Earth map."),
    ("GeoJSON", "a standard geographic feature format", "TopoJSON is converted to GeoJSON-like features for D3."),
    ("TopoJSON", "a compressed topology format sharing boundary arcs", "It keeps the world map small."),
    ("Projection", "a mathematical conversion from a globe to a flat plane", "Natural Earth converts longitude/latitude to SVG coordinates."),
    ("Responsive design", "layout that adapts to screen size", "Media queries collapse grids and reveal mobile navigation."),
    ("API", "a contract for software-to-software communication", "A future FastAPI service could deliver filtered metrics."),
    ("Time series", "ordered observations indexed by time", "Annual release counts form a simple time series."),
    ("MAE", "average absolute prediction error", "It would evaluate a future forecasting model."),
    ("Data leakage", "using information unavailable at prediction time", "It must be prevented in future temporal modeling."),
]

add_heading(doc, "30 Basic Questions", 2)
for i, (topic, definition, usage) in enumerate(concepts, 1):
    add_heading(doc, f"{i}. What is {topic}?", 3)
    add_para(doc, f"{topic} is {definition}. In this project, {usage} The main advantage is that it solves a specific, understandable problem. Its limitation is that it does not replace the other layers of the system.")

add_heading(doc, "30 Intermediate Questions", 2)
for i, (topic, definition, usage) in enumerate(concepts, 1):
    add_heading(doc, f"{i}. How is {topic} used in this project, and what alternative exists?", 3)
    add_para(doc, f"{usage} This choice keeps the MVP readable and lightweight. An alternative would be a more specialized framework or backend implementation when scale, stronger typing, collaboration, or performance becomes more important. The correct choice depends on data volume, team size, deployment constraints, and maintenance needs.")

advanced_topics = [
    ("Why is browser-side O(n) filtering acceptable?", "The dataset has only 5,850 records, so a full scan is quick. At millions of rows, use indexed server queries, pagination, caching, or columnar analytics."),
    ("How would you prevent XSS from imported data?", "Escape text before inserting it into HTML, avoid raw template interpolation, use textContent or a sanitizer, validate MIME/size, and apply a Content Security Policy."),
    ("How would you scale the importer?", "Stream or chunk parse, use a Web Worker, show progress, enforce file limits, validate schema separately, and upload to a backend for very large files."),
    ("Why is a time-ordered split required?", "Random splitting lets future patterns leak into training. Train on earlier periods and validate on later periods using rolling-origin evaluation."),
    ("How would you evaluate forecast intervals?", "Measure empirical coverage: the proportion of actual observations inside the predicted interval, plus interval width and calibration by horizon."),
    ("What is the trade-off of primary genre?", "It simplifies exclusive counts but discards multi-label information. Multi-hot encoding or fractional counting preserves more nuance."),
    ("What is the trade-off of primary country?", "It makes one marker/ranking per title but undercounts co-productions. Alternatives are fractional allocation or many-to-many tables."),
    ("How would a warehouse schema look?", "A title dimension, person dimension, genre dimension, country dimension, and fact/bridge tables for title-person, title-genre, and title-country."),
    ("How would Redis help?", "Cache expensive aggregate responses by filter key, reducing warehouse/API latency. It adds invalidation and operational complexity."),
    ("How would OAuth/SSO work?", "The frontend obtains a token from an identity provider; the API validates signatures/scopes before returning protected data."),
    ("Why use TopoJSON?", "Shared boundaries are encoded once, reducing file size and enabling a border mesh. GeoJSON is simpler but larger."),
    ("How do projections distort maps?", "Every flat projection distorts area, shape, direction, or distance. Natural Earth balances global aesthetics; equal-area projections suit quantitative choropleths."),
    ("How would you test normalizeRecord?", "Use unit cases for aliases, missing title, bad year, country arrays, comma-separated genres, numeric strings, and malicious text."),
    ("How would you test render?", "Test view routing and DOM output with a browser test runner, then visual regression screenshots for layout."),
    ("What is a medallion architecture?", "Bronze keeps raw data, Silver cleans/conforms it, and Gold stores analysis-ready aggregates. It improves traceability but adds pipeline layers."),
    ("Why use dbt?", "It manages SQL transformations, tests, lineage, and documentation in version control. It is not an ingestion scheduler."),
    ("Why use Airflow?", "It schedules and monitors task dependencies. It is powerful but operationally heavier than simple cron or managed orchestration."),
    ("How would data quality be measured?", "Uniqueness, completeness, validity, referential integrity, freshness, distribution drift, and row-count reconciliation."),
    ("What is feature drift?", "Input distributions change over time. Monitor metrics such as PSI, KS distance, category shares, and missingness."),
    ("What is concept drift?", "The relationship between inputs and target changes, causing prediction quality to deteriorate even when input distributions appear stable."),
    ("How would SHAP be used?", "It decomposes a model prediction into feature contributions. It explains the model, not necessarily the real-world causal process."),
    ("Why establish a naive baseline?", "A complex model is only useful if it beats a simple rule such as last period or seasonal last year."),
    ("When is MAPE inappropriate?", "When actual values are zero or near zero, percentage errors explode. Use MAE, RMSE, WAPE, or SMAPE."),
    ("How would anomaly detection work?", "Create expected seasonal behavior, score residuals or fit Isolation Forest, then validate alerts with domain experts."),
    ("What is caching invalidation?", "Cached answers become stale after data refresh. Use versioned keys, TTLs, or explicit invalidation after pipeline completion."),
    ("What is accessibility risk in charts?", "Color-only meaning, poor contrast, missing text alternatives, keyboard-inaccessible interactions, and unreadable responsive layouts."),
    ("How would CI/CD improve quality?", "Automated linting, tests, builds, security checks, and preview deployments catch defects before release."),
    ("How would observability work?", "Collect logs, metrics, and traces; define SLIs/SLOs for latency, errors, freshness, and availability; alert on meaningful failures."),
    ("What is reproducibility?", "A versioned combination of code, data, environment, and parameters should recreate the same output."),
    ("How would you migrate to React?", "Split cards/charts/views into components, place state in hooks/context, type records with TypeScript, and preserve pure analytics functions."),
]
add_heading(doc, "30 Advanced Questions", 2)
for i, (q, a) in enumerate(advanced_topics, 1):
    add_heading(doc, f"{i}. {q}", 3)
    add_para(doc, a + " The design decision should be justified using correctness, scale, maintainability, security, and user experience.")

project_qas = [
    ("Why was synthetic data replaced?", "The supplied real dataset made conclusions more defensible and exposed real schema constraints."),
    ("Why are there two CSV files?", "Titles contain title metadata; credits form a many-to-one relationship with titles and require a join."),
    ("Why is catalogue.js generated?", "It avoids parsing and joining multi-megabyte CSV files on every page load."),
    ("Why can the app run without a backend?", "All prepared data and logic are static files executed by the browser."),
    ("Why is localhost required?", "Serving over HTTP avoids file-origin restrictions and supports consistent loading of assets."),
    ("Why did localhost initially fail?", "The smoke-test server had stopped, so no process was listening on port 8000."),
    ("Why was the strategy card removed?", "A 3/12-width card forced content into a narrow strip and harmed readability."),
    ("Why did the first map look unrealistic?", "It used hand-drawn continent polygons as a placeholder."),
    ("How was the map corrected?", "Natural Earth TopoJSON, D3 Natural Earth projection, country paths, border mesh, and projected markers."),
    ("Why did the corrected map still fail once?", "Runtime fetching failed in the browser, so topology was loaded as a local JavaScript global."),
    ("What does filteredData return?", "Only titles satisfying all active filters and the search predicate."),
    ("How are top genres calculated?", "Group active rows by primary genre, sort counts descending, and keep the leading entries."),
    ("How is the rating donut calculated?", "Count ratings, convert counts to percentages, and create cumulative conic-gradient stops."),
    ("How are contributors joined?", "Python groups credit rows by title ID before attaching directors and cast to each title."),
    ("Why limit cast and directors?", "It controls output file size and dashboard processing, though it sacrifices completeness."),
    ("How are imported columns recognized?", "firstValue compares lowercased headers against alias lists."),
    ("What imported fields are required?", "A non-empty title and a plausible release year; type defaults to Movie when unrecognized."),
    ("Where does imported data live?", "Only in JavaScript memory for the current page session."),
    ("How do you restore the original data?", "Open Import data and choose Restore bundled data."),
    ("What happens to invalid imported rows?", "normalizeRecord returns null; filter(Boolean) removes them and the UI reports the skipped count."),
    ("How does export work?", "It serializes the current filtered subset, creates a text/csv Blob, and triggers a browser download."),
    ("What is the biggest data limitation?", "No viewership or date-added information, so business performance and true addition forecasting cannot be measured."),
    ("What is misleading in the current UI?", "Some growth/velocity comparison texts are illustrative placeholders rather than calculated historical metrics."),
    ("What is the complexity of filtering?", "O(n) for n active catalogue records per render."),
    ("What is the complexity of top-N ranking?", "Counting is O(n); sorting k categories is O(k log k)."),
    ("How is responsive design implemented?", "CSS media queries change grid spans, sidebar behavior, table columns, and button labels."),
    ("Why use local vendor files?", "The app continues working without external CDNs after setup and avoids network variability."),
    ("What would be your first production refactor?", "Escape imported text, calculate all claims, add tests, and separate analytics from HTML generation."),
    ("What would be your first data-engineering improvement?", "Create normalized bridge tables for all genres/countries/contributors rather than selecting only the first."),
    ("What would make the forecast real?", "Genuine monthly date-added data, baseline comparison, rolling validation, trained/versioned models, and monitored errors."),
]
add_heading(doc, "30 Project-Specific Questions", 2)
for i, (q, a) in enumerate(project_qas, 1):
    add_heading(doc, f"{i}. {q}", 3)
    add_para(doc, a)

viva_qas = [
    ("State the project in one sentence.", "A browser-based catalogue analytics dashboard that preprocesses real title and credit metadata and supports interactive exploration, import, and export."),
    ("What is your individual contribution?", "Translate the PRD into an MVP, integrate the supplied dataset, build preprocessing, analytics views, mapping, import/export, responsive UI, and documentation."),
    ("What is the dependent variable?", "There is no supervised learning target in the implemented system; it is descriptive analytics."),
    ("What is the unit of analysis?", "Usually one title; contributor analysis additionally uses one credit or one actor-director pair."),
    ("Why is the project not causal?", "It observes catalogue metadata without experimental assignment or confounder control."),
    ("What is the join key?", "The id field shared by titles.csv and credits.csv."),
    ("Why not use Excel?", "Excel is useful for ad-hoc analysis, but the dashboard provides repeatable definitions, richer interaction, map rendering, and reusable import/export."),
    ("Why not use pandas in the browser?", "Pandas is a Python library and would require a backend or WebAssembly runtime; JavaScript is native to browsers."),
    ("What is your strongest engineering decision?", "Precompute the title-credit join once, then keep the dashboard static and simple for the current scale."),
    ("What is your weakest current area?", "The forecast and some KPI comparisons are illustrative, not data-derived production analytics."),
    ("How do you ensure correctness?", "Schema checks, row-count reconciliation, syntax checks, HTTP asset checks, explicit missing-value rules, and future automated tests."),
    ("What happens with a malformed file?", "The importer catches parse errors, rejects invalid rows, and displays an error without replacing the active catalogue."),
    ("What is one security risk?", "Unescaped imported text inserted into HTML could permit XSS; production code should use textContent/sanitization."),
    ("What is one performance risk?", "Large imports cause repeated main-thread O(n) filtering and DOM rendering."),
    ("What is one ethical consideration?", "Avoid claiming viewer preferences, demand, or performance from metadata alone."),
    ("How would you make results reproducible?", "Version the source archive, preprocessing code, generated catalogue, dependencies, and deployment commit."),
    ("How would you validate region mapping?", "Use ISO country libraries, unit tests, and reconciliation of unmapped/Unknown codes."),
    ("Why Natural Earth?", "It provides lightweight, recognized global boundary data appropriate for a dashboard overview."),
    ("What would you demonstrate live?", "Import a small CSV, filter it, inspect the geographic and library views, export the subset, then restore bundled data."),
    ("What is the final conclusion?", "The MVP successfully makes catalogue metadata explorable, but reliable strategic prediction requires richer data and a real validated backend/ML pipeline."),
]
add_heading(doc, "20 Viva Questions", 2)
for i, (q, a) in enumerate(viva_qas, 1):
    add_heading(doc, f"{i}. {q}", 3)
    add_para(doc, a)

# Part IX
section_intro(doc, "EXPLAIN THIS PROJECT TO AN INTERVIEWER", "Ready-to-use versions for different time limits.")
scripts = [
    ("30-second version", "I built a Netflix catalogue analytics dashboard using HTML, CSS, JavaScript, Python, D3, and TopoJSON. I processed 5,850 titles and 77,801 cast/director credits, joined and normalized them into a browser-ready dataset, and created interactive overview, trend, geographic, contributor, and library views. The app also supports filtering, search, CSV export, and importing another CSV or JSON dataset. It is a descriptive MVP; the forecasting screen is illustrative, not a trained production model."),
    ("1-minute version", "This project converts raw Netflix-like metadata into an interactive decision-support dashboard. The source has separate title and credit tables, so I wrote a Python preprocessing pipeline to parse list fields, normalize types and countries, group actors/directors by title ID, and emit a compact browser dataset. The frontend is dependency-light: semantic HTML, responsive CSS, and JavaScript for state, filtering, aggregation, SVG charts, import, and export. D3 and TopoJSON render a real Natural Earth world map. I made deliberate corrections when the real dataset lacked date-added and availability fields, so the final analysis uses release year and production origin. For production, I would add a warehouse/API, tests, security hardening, and real time-series backtesting."),
    ("3-minute version", "The original PRD described a large 20-week platform with pipelines, a warehouse, APIs, machine learning, and reporting. I scoped that into a working frontend MVP while preserving the main analytical user journey. The supplied archive contains 5,850 titles and 77,801 contributor credits. I first inspected the schema and discovered it differs from the PRD dataset: it has release year, runtime, seasons, IMDb/TMDb metrics, genres, production countries, actors, and directors, but no Netflix date-added or true regional availability. I therefore changed the dashboard language and avoided pretending those fields existed. A Python ETL script reads both CSVs, safely parses list-like values, creates a dictionary keyed by title ID, joins directors and cast in near-linear time, derives country/region and primary genre, and generates catalogue.js. In the browser, app.js stores filter state and recomputes descriptive analytics using filter, reduce, grouping, sorting, and top-N selection. Reusable rendering functions create KPI cards, line and bar charts, donut charts, heatmaps, tables, and a Natural Earth map rendered with D3 and TopoJSON. I also added CSV/JSON import with common column aliases, row validation, skipped-row reporting, and restore-original-data behavior. The major limitation is that the forecast is illustrative. A defensible ML version would need genuine monthly addition or demand data, time-based validation, naive baselines, Prophet/SARIMA comparison, and monitored errors."),
    ("5-minute version", "I started from a product requirements document for a Netflix Content Analysis Platform. The long-term vision included a medallion data pipeline, a warehouse, FastAPI, forecasting, anomaly detection, reporting, and an API. Rather than claim all of that was built, I delivered a self-contained analytical MVP and documented the production path. The data engineering problem was important because titles and contributors are separate. I used Python's csv module and ast.literal_eval, grouped credits in a defaultdict keyed by title ID, limited contributor arrays for browser size, normalized categorical and numeric fields, derived primary country/region and genre, and serialized 5,850 normalized records. The frontend separates structure, presentation, and behavior into index.html, styles.css, and app.js. A central state object controls view, type, region, genre, year, and search. filteredData performs an O(n) scan, which is acceptable at this scale. group uses reduce for one-pass frequency distributions, topEntries ranks categories, and individual view functions calculate the values needed for each screen. Charts are implemented with SVG or CSS, while the map uses a Natural Earth projection and TopoJSON boundaries. During development I corrected several real issues: Unicode extraction, workspace permissions, the temporary server stopping, a narrow card breaking layout, an unrealistic placeholder map, and browser fetch restrictions. The importer now accepts CSV or JSON, recognizes common aliases, validates title and year, normalizes each row into the same canonical schema, and rerenders every view. For production I would address XSS escaping, automated tests, all-country markers, calculated KPI comparisons, large-file workers, a normalized database schema, APIs, authentication, caching, observability, and real forecasting with rolling backtests. The strongest lesson is scope honesty: this is a good descriptive analytics product, but metadata alone cannot prove audience demand or content performance."),
]
for title, text in scripts:
    add_heading(doc, title, 2)
    add_para(doc, text, size=10.8)

add_heading(doc, "Final Defense Checklist", 2)
for item in [
    "Say exactly which features are implemented.",
    "Do not describe the illustrative projection as trained AI.",
    "Know the two source tables and their join key.",
    "Be able to explain filter, reduce, dictionary join, and Natural Earth projection.",
    "State at least three limitations without being prompted.",
    "Describe the first production improvements: security, testing, calculated KPIs, backend, and real ML validation.",
    "Demonstrate import, filtering, map, export, and restore flow.",
]:
    add_bullet(doc, item)

add_heading(doc, "Glossary", 2)
glossary = [
    ("MVP", "Minimum viable product: smallest useful version that tests the core experience."),
    ("PRD", "Product requirements document."),
    ("EDA", "Exploratory data analysis."),
    ("ETL", "Extract, transform, load."),
    ("DOM", "Document Object Model."),
    ("SVG", "Scalable Vector Graphics."),
    ("API", "Application Programming Interface."),
    ("XSS", "Cross-site scripting."),
    ("SSO", "Single sign-on."),
    ("CI/CD", "Continuous integration and continuous delivery/deployment."),
    ("SLA/SLO/SLI", "Service commitment, objective, and measured indicator."),
    ("MAPE/MAE/RMSE", "Common forecast error metrics."),
]
add_table(doc, ["Term", "Meaning"], glossary, [1.4, 5.1])

# Final page
page_break(doc)
add_para(doc, "END OF HANDBOOK", size=10, color=RED, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=150, after=16)
add_para(doc, "You should now be able to understand the project, demonstrate it, explain its engineering decisions, defend its limitations, and describe a credible path from MVP to production.", size=14, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
add_para(doc, "Prepared for Kshitij", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.core_properties.title = "Netflix Content Analysis - Complete Project Handbook"
doc.core_properties.subject = "Beginner-to-advanced technical and interview handbook"
doc.core_properties.author = "Kshitij"
doc.core_properties.keywords = "Netflix, content analysis, dashboard, JavaScript, Python, D3, interview, handbook"
doc.save(OUT)
print(OUT)
